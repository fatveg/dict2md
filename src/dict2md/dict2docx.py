import docx
from .dict2md import title_clean
from typing import Dict, Any, List, Optional


def dict2docx(
    d: Dict[str, Any],
    *,
    title_keys: Optional[List[str]] = None,
    output_filename: str = "output.docx",
    _document_object: Optional[docx.Document] = None,
    _heading_level: int = 1,
) -> None:
    """
    Convert a dictionary to a Word document (.docx).

    :param d: The dictionary to convert.
    :param title_keys: A list of keys to include at the top level.
    :param heading_level: The current heading level (default=1).
    :param output_filename: The filename of the output Word document (default='output.docx').
    :return: None
    """

    if _document_object is None:
        _document_object = docx.Document()

    for key, value in d.items():
        if title_keys and key.casefold() in (k.casefold() for k in title_keys):
            _document_object.add_heading(value, 0)
        elif isinstance(value, str):
            _document_object.add_heading(title_clean(key), _heading_level)
            _document_object.add_paragraph(value)
        elif isinstance(value, list):
            _document_object.add_heading(title_clean(key), _heading_level)
            for item in value:
                _document_object.add_paragraph(item, style="List Bullet")
        elif isinstance(value, dict):
            _document_object.add_heading(title_clean(key), _heading_level)
            dict2docx(
                value,
                title_keys=title_keys,
                _heading_level=_heading_level + 1,
                _document_object=_document_object,
            )
        else:
            _document_object.add_heading(title_clean(key), _heading_level)
            _document_object.add_paragraph(str(value))

    # Save the Word document only if we are at the top level
    if _heading_level == 1:
        _document_object.save(output_filename)
